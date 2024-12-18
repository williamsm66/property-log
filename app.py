from flask import Flask, render_template, request, jsonify
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
from property_scraper import PropertyScraper
import asyncio
import os
import json
import uuid
import shutil
import tempfile
import logging
import anthropic
import tiktoken
import pytesseract
from pdf2image import convert_from_path
import io
import subprocess
from docx import Document
import PyPDF2
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
from werkzeug.serving import WSGIRequestHandler
import time
from pathlib import Path
from google.cloud import vision
from google.cloud import documentai_v1 as documentai
from datetime import timedelta  # Import timedelta for viewing schedule
import gc  # Import garbage collector
from threading import Thread
import psutil
import zipfile
from gevent import monkey; monkey.patch_all()

# Initialize Flask app first
app = Flask(__name__)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),  # Output to console
        logging.FileHandler('app.log')  # Also save to file
    ]
)
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

# Ensure Flask's built-in logger also shows INFO messages
logging.getLogger('werkzeug').setLevel(logging.INFO)

# Load environment variables
logger.info("Starting application initialization...")
load_dotenv()
logger.info("Environment variables loaded")

# Initialize Google Cloud clients
vision_client = None
try:
    vision_client = vision.ImageAnnotatorClient()
    logger.info("Google Vision client initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize Google Vision client: {str(e)}")

# Check required environment variables
required_vars = ['CLAUDE_API_KEY', 'DATABASE_URL', 'GOOGLE_APPLICATION_CREDENTIALS', 'GOOGLE_CLOUD_PROJECT']
missing_vars = [var for var in required_vars if not os.getenv(var)]
if missing_vars:
    logger.warning(f"Missing required environment variables: {', '.join(missing_vars)}")
else:
    logger.info("All required environment variables are set")

logger.info("Application initialization completed")

# Configure database and uploads
basedir = os.path.abspath(os.path.dirname(__file__))

# Use PostgreSQL in production, SQLite in development
if os.environ.get('DATABASE_URL'):
    # Handle Render.com's DATABASE_URL format
    database_url = os.environ.get('DATABASE_URL')
    if database_url.startswith('postgres://'):
        database_url = database_url.replace('postgres://', 'postgresql://', 1)
    app.config['SQLALCHEMY_DATABASE_URI'] = database_url
else:
    db_path = os.path.join(basedir, 'properties.db')
    app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{db_path}'

app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = os.path.join(basedir, 'uploads')
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Create a directory for storing documents
STORAGE_DIR = Path(os.path.dirname(os.path.abspath(__file__))) / 'document_storage'
try:
    STORAGE_DIR.mkdir(mode=0o777, parents=True, exist_ok=True)
    # Ensure parent directory has right permissions too
    STORAGE_DIR.parent.chmod(0o777)
    STORAGE_DIR.chmod(0o777)
except Exception as e:
    logger.error(f"Error creating storage directory: {str(e)}")

# Set database file permissions if it doesn't exist
if not os.path.exists(db_path):
    open(db_path, 'a').close()  # Create file if it doesn't exist
    os.chmod(db_path, 0o666)  # Set read/write permissions for user and group

# Initialize database
db = SQLAlchemy(app)

class Property(db.Model):
    __tablename__ = 'property'
    
    id = db.Column(db.Integer, primary_key=True)
    rightmove_url = db.Column(db.String(500), nullable=True)
    initial_cash = db.Column(db.Float)
    purchase_price = db.Column(db.Float)
    rooms = db.Column(db.Integer)
    monthly_rent = db.Column(db.Float)
    valuation_after = db.Column(db.Float)
    renovation_cost = db.Column(db.Float)
    bridging_duration = db.Column(db.Integer)
    void_period = db.Column(db.Integer)
    mortgage_ltv = db.Column(db.Float)
    mortgage_rate = db.Column(db.Float)
    lender_fee = db.Column(db.Float)
    bridging_rate = db.Column(db.Float)
    arrangement_rate = db.Column(db.Float)
    broker_rate = db.Column(db.Float)
    management_fee = db.Column(db.Float)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    main_photo = db.Column(db.String(500), nullable=True)
    floorplan = db.Column(db.String(500), nullable=True)
    description = db.Column(db.Text, nullable=True)
    key_features = db.Column(db.Text, nullable=True)
    address = db.Column(db.String(500), nullable=True)
    is_auction = db.Column(db.Boolean, default=False)
    estate_agent = db.Column(db.String(200), nullable=True)
    nearest_station = db.Column(db.String(200), nullable=True)
    station_distance = db.Column(db.Float, nullable=True)
    bedrooms = db.Column(db.Integer, nullable=True)
    bathrooms = db.Column(db.Integer, nullable=True)
    property_type = db.Column(db.String(100), nullable=True)
    
    # New fields
    legal_pack_url = db.Column(db.String(500), nullable=True)
    legal_pack_available = db.Column(db.Boolean, default=False)
    risk_level = db.Column(db.String(20), nullable=True)
    key_risks = db.Column(db.Text, nullable=True)
    extra_fees = db.Column(db.Float, default=0)
    auction_date = db.Column(db.Date, nullable=True)
    
    # Legal pack fields
    legal_pack_analysis = db.Column(db.Text, nullable=True)
    legal_pack_qa_history = db.Column(db.Text, nullable=True)  # Store Q&A history as JSON
    legal_pack_documents = db.Column(db.Text, nullable=True)  # Store documents content as JSON
    legal_pack_summary_pdf = db.Column(db.String(500), nullable=True)  # Path to the PDF summary
    legal_pack_analyzed_at = db.Column(db.DateTime, nullable=True)
    legal_pack_session_id = db.Column(db.String(100), nullable=True)  # To link with legal doc analyzer session
    
    # Viewing dates fields
    viewing_date_1 = db.Column(db.DateTime, nullable=True)
    viewing_date_2 = db.Column(db.DateTime, nullable=True)
    viewing_date_3 = db.Column(db.DateTime, nullable=True)
    viewing_date_4 = db.Column(db.DateTime, nullable=True)
    
    # Calculated fields
    stamp_duty = db.Column(db.Float, nullable=True)
    total_purchase_fees = db.Column(db.Float, nullable=True)
    total_money_needed = db.Column(db.Float, nullable=True)
    cash_left_in_deal = db.Column(db.Float, nullable=True)
    annual_profit = db.Column(db.Float, nullable=True)
    total_roi = db.Column(db.Float, nullable=True)
    total_yield = db.Column(db.Float, nullable=True)

    def to_dict(self):
        return {
            'id': self.id,
            'rightmove_url': self.rightmove_url,
            'initial_cash': self.initial_cash,
            'purchase_price': self.purchase_price,
            'rooms': self.rooms,
            'monthly_rent': self.monthly_rent,
            'valuation_after': self.valuation_after,
            'renovation_cost': self.renovation_cost,
            'bridging_duration': self.bridging_duration,
            'void_period': self.void_period,
            'mortgage_ltv': self.mortgage_ltv,
            'mortgage_rate': self.mortgage_rate,
            'lender_fee': self.lender_fee,
            'bridging_rate': self.bridging_rate,
            'arrangement_rate': self.arrangement_rate,
            'broker_rate': self.broker_rate,
            'management_fee': self.management_fee,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'main_photo': self.main_photo,
            'floorplan': self.floorplan,
            'description': self.description,
            'key_features': json.loads(self.key_features) if self.key_features else [],
            'address': self.address,
            'is_auction': self.is_auction,
            'estate_agent': self.estate_agent,
            'nearest_station': self.nearest_station,
            'station_distance': self.station_distance,
            'bedrooms': self.bedrooms,
            'bathrooms': self.bathrooms,
            'property_type': self.property_type,
            'legal_pack_url': self.legal_pack_url,
            'legal_pack_available': self.legal_pack_available,
            'risk_level': self.risk_level,
            'key_risks': self.key_risks,
            'extra_fees': self.extra_fees,
            'auction_date': self.auction_date.isoformat() if self.auction_date else None,
            'stamp_duty': self.stamp_duty,
            'total_purchase_fees': self.total_purchase_fees,
            'total_money_needed': self.total_money_needed,
            'cash_left_in_deal': self.cash_left_in_deal,
            'annual_profit': self.annual_profit,
            'total_roi': self.total_roi,
            'total_yield': self.total_yield,
            'legal_pack_analysis': self.legal_pack_analysis,
            'legal_pack_qa_history': json.loads(self.legal_pack_qa_history) if self.legal_pack_qa_history else [],
            'legal_pack_documents': json.loads(self.legal_pack_documents) if self.legal_pack_documents else [],
            'legal_pack_summary_pdf': self.legal_pack_summary_pdf,
            'legal_pack_analyzed_at': self.legal_pack_analyzed_at.isoformat() if self.legal_pack_analyzed_at else None,
            'legal_pack_session_id': self.legal_pack_session_id,
            'viewing_date_1': self.viewing_date_1.isoformat() if self.viewing_date_1 else None,
            'viewing_date_2': self.viewing_date_2.isoformat() if self.viewing_date_2 else None,
            'viewing_date_3': self.viewing_date_3.isoformat() if self.viewing_date_3 else None,
            'viewing_date_4': self.viewing_date_4.isoformat() if self.viewing_date_4 else None,
        }

class Analysis(db.Model):
    __tablename__ = 'analysis'
    id = db.Column(db.Integer, primary_key=True)
    property_id = db.Column(db.Integer, db.ForeignKey('property.id'))
    content = db.Column(db.Text)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)

    def __repr__(self):
        return f'<Analysis {self.id} for property {self.property_id}>'

def init_db():
    """Initialize database and create tables."""
    try:
        app.logger.info("Starting database initialization...")
        with app.app_context():
            db.create_all()
            app.logger.info("Database tables created successfully")
    except Exception as e:
        app.logger.error(f"Failed to initialize database: {str(e)}")
        raise

# Initialize database before running the app
init_db()
app.logger.info("Database initialization completed")

# Add new model for document sessions
class DocumentSession(db.Model):
    __tablename__ = 'document_sessions'
    id = db.Column(db.String(100), primary_key=True)
    documents = db.Column(db.JSON)
    initial_analysis = db.Column(db.Text)
    qa_history = db.Column(db.JSON)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    property_id = db.Column(db.Integer, db.ForeignKey('property.id'))

class Document(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(255), nullable=False)
    status = db.Column(db.String(50), default='pending')  # pending, processing, completed, failed
    error = db.Column(db.Text)
    text_content = db.Column(db.Text)
    processed_pages = db.Column(db.Integer, default=0)
    total_pages = db.Column(db.Integer)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

def count_tokens(text):
    """Count tokens in text using tiktoken."""
    try:
        # Use cl100k_base encoding which is used by Claude
        encoding = tiktoken.get_encoding("cl100k_base")
        # Convert text to string if it isn't already
        text_str = str(text)
        # Process text in chunks to avoid recursion
        chunk_size = 100000  # Process 100KB at a time
        total_tokens = 0
        for i in range(0, len(text_str), chunk_size):
            chunk = text_str[i:i + chunk_size]
            total_tokens += len(encoding.encode(chunk))
        return total_tokens
    except Exception as e:
        logger.error(f"Error counting tokens: {str(e)}")
        # Return a conservative estimate if tiktoken fails
        return len(str(text).split()) * 2

def save_documents(session_id, processed_files, initial_analysis=None, qa_history=None):
    """Save documents and analysis history to database."""
    try:
        session = DocumentSession.query.get(session_id)
        if not session:
            session = DocumentSession(
                id=session_id,
                documents=processed_files,
                initial_analysis=initial_analysis,
                qa_history=qa_history or []
            )
            db.session.add(session)
        else:
            session.documents = processed_files
            if initial_analysis is not None:
                session.initial_analysis = initial_analysis
            if qa_history is not None:
                session.qa_history = qa_history
        
        db.session.commit()
        return True
    except Exception as e:
        logger.error(f"Error saving documents to database: {str(e)}")
        return False

def load_documents(session_id):
    """Load documents and analysis history from database."""
    try:
        session = DocumentSession.query.get(session_id)
        if not session:
            logger.error(f"Session not found in database: {session_id}")
            return None, None, None
            
        return session.documents, session.initial_analysis, session.qa_history
    except Exception as e:
        logger.error(f"Error loading documents from database: {str(e)}")
        return None, None, None

def extract_text_from_pdf(file_path):
    """Extract text from PDF file."""
    app.logger.info(f"Starting PDF extraction for: {file_path}")
    text_content = []
    
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            total_pages = len(pdf_reader.pages)
            app.logger.info(f"PDF has {total_pages} pages")
            
            # Process in batches of 10 pages
            BATCH_SIZE = 10
            MAX_PAGES = 50  # Maximum total pages to process
            
            for batch_start in range(0, total_pages, BATCH_SIZE):
                try:
                    end_page = min(batch_start + BATCH_SIZE, total_pages, MAX_PAGES)
                    batch_text = []
                    
                    for page_num in range(batch_start, end_page):
                        app.logger.info(f"Processing page {page_num+1}/{total_pages}")
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text().strip()
                        
                        if not page_text or len(page_text) < 100:
                            app.logger.info(f"Page {page_num+1} has insufficient text ({len(page_text)} chars), attempting OCR")
                            try:
                                images = convert_from_path(file_path, first_page=page_num+1, last_page=page_num+1)
                                if images:
                                    app.logger.info(f"Successfully converted page {page_num+1} to image, starting OCR")
                                    page_text = process_scanned_page(images[0])
                                    app.logger.info(f"OCR completed for page {page_num+1}, extracted {len(page_text)} chars")
                                    del images
                                else:
                                    app.logger.warning(f"No images extracted from page {page_num+1}")
                            except Exception as e:
                                app.logger.error(f"Error during OCR for page {page_num+1}: {str(e)}")
                        else:
                            app.logger.info(f"Successfully extracted {len(page_text)} chars from page {page_num+1} using PyPDF2")
                        
                        batch_text.append(page_text)
                        gc.collect()
                        app.logger.info(f"Memory usage after page {page_num+1}: {psutil.Process().memory_info().rss / 1024 / 1024:.2f} MB")
                    
                    text_content.extend(batch_text)
                    gc.collect()
                    app.logger.info(f"Completed batch {batch_start+1}-{end_page}, total content length: {sum(len(t) for t in batch_text)}")
                    
                except Exception as e:
                    app.logger.error(f"Error processing pages {batch_start+1}-{end_page}: {str(e)}")
                    continue
        
        if total_pages > MAX_PAGES:
            app.logger.warning(f"PDF has {total_pages} pages, but only processed first {MAX_PAGES} pages")
            text_content.append(f"\n[Note: Only the first {MAX_PAGES} pages were processed due to size limits]")
        
        final_text = "\n".join(text_content)
        app.logger.info(f"PDF extraction completed. Total content length: {len(final_text)}")
        return final_text
            
    except Exception as e:
        app.logger.error(f"Error in PDF extraction: {str(e)}")
        raise

def init_vision_client():
    """Initialize or reinitialize the Vision client."""
    global vision_client
    try:
        if not vision_client:
            vision_client = vision.ImageAnnotatorClient()
            logger.info("Google Vision client initialized successfully")
        return True
    except Exception as e:
        logger.error(f"Failed to initialize Vision client: {str(e)}")
        return False

def process_scanned_page(image):
    """Process a scanned page using OCR."""
    global vision_client
    try:
        # Try to initialize Vision client if it's not already initialized
        if not vision_client and not init_vision_client():
            raise Exception("Vision client not available")
            
        # Convert PIL image to bytes
        with io.BytesIO() as bio:
            image.save(bio, format='PNG')
            image_bytes = bio.getvalue()
        
        # Create Vision API image
        vision_image = vision.Image(content=image_bytes)
        
        # Perform OCR
        response = vision_client.text_detection(image=vision_image)
        if response.error.message:
            raise Exception(response.error.message)
            
        texts = response.text_annotations
        if texts:
            return texts[0].description
        return ""
        
    except Exception as e:
        logger.error(f"OCR Error: {str(e)}")
        # Try to reinitialize client for next time
        init_vision_client()
        return ""
    finally:
        # Clean up
        if 'image_bytes' in locals():
            del image_bytes
        gc.collect()

def extract_text_from_doc(doc_path):
    """Extract text from Word documents using multiple methods for maximum compatibility."""
    try:
        logger.info(f"Attempting to extract text from Word document: {doc_path}")
        
        # Try docx2txt first as it's the most reliable for both .doc and .docx
        try:
            logger.info("Attempting to extract text using docx2txt")
            import docx2txt
            text = docx2txt.process(doc_path)
            if text and text.strip():
                logger.info(f"Successfully extracted {len(text)} characters using docx2txt")
                return text
        except Exception as e:
            logger.warning(f"docx2txt extraction failed: {str(e)}")

        # If docx2txt fails and it's a .docx file, try python-docx as fallback
        if doc_path.lower().endswith('.docx'):
            try:
                logger.info("Falling back to python-docx for .docx file")
                doc = Document(doc_path)
                full_text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                full_text.append(cell.text)
                
                text = '\n'.join(full_text)
                if text.strip():
                    logger.info(f"Successfully extracted {len(text)} characters using python-docx")
                    return text
            except Exception as e:
                logger.warning(f"python-docx extraction failed: {str(e)}")
        
        # If all else fails, try a simple binary read
        try:
            logger.info("Attempting raw text extraction as last resort")
            with open(doc_path, 'rb') as file:
                raw_content = file.read()
                # Try to decode with different encodings
                for encoding in ['utf-8', 'latin-1', 'cp1252']:
                    try:
                        text = raw_content.decode(encoding)
                        # Clean the text by removing non-printable characters
                        text = ''.join(char for char in text if char.isprintable() or char in '\n\r\t')
                        if text.strip():
                            logger.info(f"Successfully extracted text using raw binary read with {encoding} encoding")
                            return text
                    except:
                        continue
        except Exception as e:
            logger.warning(f"Raw text extraction failed: {str(e)}")

        logger.error("All text extraction methods failed")
        return None
            
    except Exception as e:
        logger.error(f"Error in extract_text_from_doc: {str(e)}")
        return None

def process_document(file_path):
    """Process a single document and return its text content."""
    try:
        _, ext = os.path.splitext(file_path.lower())
        logger.info(f"Processing document: {file_path} (type: {ext})")
        
        if ext == '.pdf':
            logger.info(f"Extracting text from PDF: {file_path}")
            content = extract_text_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            logger.info(f"Extracting text from Word document: {file_path}")
            content = extract_text_from_doc(file_path)
        elif ext == '.zip':
            logger.info(f"Processing ZIP file: {file_path}")
            processed_files, failed_files, processing_summary = process_zip_file(file_path)
            if processed_files:
                content = "\n\n".join(doc['content'] for doc in processed_files)
                logger.info(f"Successfully extracted content from {len(processed_files)} files in ZIP")
            else:
                logger.warning("No content extracted from ZIP file")
                content = None
        else:
            logger.warning(f"Unsupported file type: {ext} for file {file_path}")
            return None
            
        if content:
            logger.info(f"Successfully extracted {len(content)} characters from {file_path}")
        else:
            logger.warning(f"No content extracted from {file_path}")
        return content
    except Exception as e:
        logger.error(f"Error processing file {file_path}: {str(e)}")
        return None

def process_zip_file(zip_file_path):
    """Process a ZIP file and extract its contents."""
    logger.info(f"Starting to process ZIP file: {zip_file_path}")
    processed_files = []
    failed_files = []
    processing_summary = []
    total_tokens = 0
    
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            logger.info(f"Created temporary directory: {temp_dir}")
            with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
                # Log ZIP contents
                files_in_zip = [f for f in zip_ref.namelist() if not f.startswith('__MACOSX/') and not f.startswith('._')]
                logger.info(f"Files in ZIP: {files_in_zip}")
                
                # Extract files while filtering out macOS metadata
                for file_info in zip_ref.filelist:
                    if not file_info.filename.startswith('__MACOSX/') and not file_info.filename.startswith('._'):
                        zip_ref.extract(file_info, temp_dir)
                        logger.info(f"Extracted: {file_info.filename}")
                
                # Process each file in the zip
                for root, _, files in os.walk(temp_dir):
                    for file in sorted(files):  # Sort files to ensure consistent processing order
                        if file.startswith('.') or file.startswith('~'):  # Skip hidden and temporary files
                            logger.info(f"Skipping hidden/temp file: {file}")
                            continue
                            
                        file_path = os.path.join(root, file)
                        logger.info(f"Processing file from ZIP: {file}")
                        try:
                            content = process_document(file_path)
                            if content and content.strip():
                                num_tokens = count_tokens(content)
                                total_tokens += num_tokens
                                processed_files.append({
                                    'name': file,
                                    'content': content,
                                    'length': len(content),
                                    'tokens': num_tokens
                                })
                                msg = f"Successfully processed {file} ({num_tokens} tokens)"
                                logger.info(msg)
                                processing_summary.append(msg)
                            else:
                                msg = f"Failed to extract content from {file}"
                                logger.warning(msg)
                                failed_files.append(file)
                                processing_summary.append(msg)
                        except Exception as e:
                            msg = f"Error processing {file}: {str(e)}"
                            logger.error(msg)
                            failed_files.append(file)
                            processing_summary.append(msg)

    except Exception as e:
        logger.error(f"Error processing ZIP file: {str(e)}")
        raise
    # Save processing results
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    results = {
        "documents": processed_files,
        "failed_files": failed_files,
        "processed_at": datetime.now().isoformat(),
        "total_documents": len(processed_files) + len(failed_files),
        "total_tokens": total_tokens,
        "processing_summary": processing_summary
    }
    
    results_filename = f"processing_results_{timestamp}.json"
    results_filepath = STORAGE_DIR / results_filename
    
    with open(results_filepath, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    
    logger.info(f"Processing results saved to {results_filepath}")
    logger.info(f"Total tokens across all documents: {total_tokens}")
    
    return processed_files, failed_files, "\n".join(processing_summary)

def analyze_with_claude(documents_content, processing_summary=None, follow_up_question=None, initial_analysis=None, qa_history=None):
    """Analyze all documents together using Claude API."""
    try:
        # Initialize the Anthropic client with the API key from environment variables
        api_key = os.getenv('CLAUDE_API_KEY')
        if not api_key:
            logger.error("CLAUDE_API_KEY environment variable is not set")
            raise ValueError("CLAUDE_API_KEY environment variable is not set")
        
        logger.info(f"Analyzing documents with Claude (follow_up: {'yes' if follow_up_question else 'no'})")
        
        try:
            client = anthropic.Anthropic(api_key=api_key, timeout=300)  # 5 minute timeout
            logger.info("Successfully initialized Anthropic client")
        except Exception as client_error:
            logger.error(f"Failed to initialize Anthropic client: {str(client_error)}")
            raise
        
        # Process documents in smaller batches
        MAX_BATCH_TOKENS = 12000  # Leave room for prompt and response
        current_batch = []
        current_batch_tokens = 0
        all_responses = []
        
        for doc in documents_content:
            # Calculate tokens for this document
            doc_text = f"\n{'='*50}\nDOCUMENT: {doc['name']}\n{'='*50}\n{doc['content']}"
            doc_tokens = count_tokens(doc_text)
            
            if current_batch_tokens + doc_tokens > MAX_BATCH_TOKENS:
                # Process current batch
                batch_response = process_document_batch(current_batch, client)
                if batch_response:
                    all_responses.append(batch_response)
                # Clear batch and memory
                current_batch = []
                current_batch_tokens = 0
                
                # Force garbage collection
                gc.collect()
            
            current_batch.append(doc_text)
            current_batch_tokens += doc_tokens
            # Clear individual document text
            doc_text = None
            
        # Process final batch if any
        if current_batch:
            batch_response = process_document_batch(current_batch, client)
            if batch_response:
                all_responses.append(batch_response)
        
        # Combine all responses
        combined_analysis = "\n\n".join(all_responses)
        
        # Final analysis of combined results
        if not follow_up_question:
            system_prompt = """You are an expert conveyancer. Review and consolidate the following analyses of legal pack documents into a single, coherent summary. Focus on the most important findings and risks."""
            
            try:
                final_response = client.messages.create(
                    model="claude-3-sonnet-20240229",
                    max_tokens=4096,
                    system=system_prompt,
                    messages=[{
                        "role": "user",
                        "content": f"Previous analyses:\n\n{combined_analysis}\n\nProvide a consolidated summary focusing on:\n1. Most significant risks and findings\n2. Key recommendations\n3. Critical missing information"
                    }],
                    temperature=0
                )
                return final_response.content[0].text
            except Exception as api_error:
                logger.error(f"Claude API error during final analysis: {str(api_error)}")
                raise ValueError(f"Failed to get final analysis from Claude API: {str(api_error)}")
        else:
            # For follow-up questions, use the original logic
            system_prompt = """You are an expert conveyancer analyzing a legal pack for an auction property. You have previously provided a comprehensive analysis, and now need to answer a specific follow-up question."""
            
            context = "Here is the initial analysis of the legal pack:\n\n"
            context += initial_analysis + "\n\n"
            
            if qa_history:
                context += "Previous questions and answers:\n\n"
                for qa in qa_history:
                    context += f"Q: {qa['question']}\nA: {qa['answer']}\n\n"
            
            try:
                response = client.messages.create(
                    model="claude-3-sonnet-20240229",
                    max_tokens=4096,
                    system=system_prompt,
                    messages=[{
                        "role": "user",
                        "content": f"Context:\n{context}\n\nDocuments analyzed:\n{combined_analysis}\n\nQuestion: {follow_up_question}"
                    }],
                    temperature=0
                )
                return response.content[0].text
            except Exception as api_error:
                logger.error(f"Claude API error during follow-up: {str(api_error)}")
                raise ValueError(f"Failed to get answer from Claude API: {str(api_error)}")
                
    except Exception as e:
        logger.error(f"Error in analyze_with_claude: {str(e)}")
        raise
    finally:
        # Clean up
        gc.collect()

def process_document_batch(document_batch, client):
    """Process a batch of documents with Claude."""
    try:
        system_prompt = """You are an expert conveyancer analyzing a legal pack for an auction property. Analyze this batch of documents and identify key risks and important information."""
        
        user_prompt = """Analyze these legal pack documents, focusing on:
1. Key risks and issues
2. Legal restrictions or covenants
3. Financial obligations
4. Development constraints
5. Environmental concerns
6. Missing critical information

Format your response in these sections:
1. DOCUMENT SUMMARY
2. KEY FINDINGS AND RISKS
3. IMPORTANT INFORMATION"""

        response = client.messages.create(
            model="claude-3-sonnet-20240229",
            max_tokens=4096,
            system=system_prompt,
            messages=[{
                "role": "user",
                "content": f"Documents to analyze:\n{''.join(document_batch)}\n\n{user_prompt}"
            }],
            temperature=0
        )
        
        return response.content[0].text if response else None
    except Exception as e:
        logger.error(f"Error processing document batch: {str(e)}")
        return None
    finally:
        # Clean up batch memory
        document_batch.clear()
        gc.collect()

def process_document_async(doc_id):
    """Process document in background."""
    try:
        document = Document.query.get(doc_id)
        if not document:
            return
        
        document.status = 'processing'
        db.session.commit()
        
        # Get file path from document
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], document.filename)
        
        # Process in chunks of 10 pages
        CHUNK_SIZE = 10
        text_chunks = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            total_pages = len(pdf_reader.pages)
            document.total_pages = total_pages
            db.session.commit()
            
            for start_page in range(0, total_pages, CHUNK_SIZE):
                try:
                    end_page = min(start_page + CHUNK_SIZE, total_pages)
                    chunk_text = []
                    
                    for page_num in range(start_page, end_page):
                        app.logger.info(f"Processing page {page_num+1}/{total_pages}")
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text().strip()
                        
                        if not page_text or len(page_text) < 100:
                            app.logger.info(f"Page {page_num+1} has insufficient text ({len(page_text)} chars), attempting OCR")
                            try:
                                images = convert_from_path(file_path, first_page=page_num+1, last_page=page_num+1)
                                if images:
                                    app.logger.info(f"Successfully converted page {page_num+1} to image, starting OCR")
                                    page_text = process_scanned_page(images[0])
                                    app.logger.info(f"OCR completed for page {page_num+1}, extracted {len(page_text)} chars")
                                    del images
                                else:
                                    app.logger.warning(f"No images extracted from page {page_num+1}")
                            except Exception as e:
                                app.logger.error(f"Error during OCR for page {page_num+1}: {str(e)}")
                        else:
                            app.logger.info(f"Successfully extracted {len(page_text)} chars from page {page_num+1} using PyPDF2")
                        
                        chunk_text.append(page_text)
                        document.processed_pages = page_num + 1
                        db.session.commit()
                        gc.collect()
                    
                    text_chunks.extend(chunk_text)
                    gc.collect()
                    app.logger.info(f"Completed batch {start_page+1}-{end_page}, total content length: {sum(len(t) for t in chunk_text)}")
                    
                except Exception as e:
                    app.logger.error(f"Error processing pages {start_page+1}-{end_page}: {str(e)}")
                    continue
        
        document.text_content = "\n".join(text_chunks)
        document.status = 'completed'
        db.session.commit()
        
    except Exception as e:
        app.logger.error(f"Error processing document {doc_id}: {str(e)}")
        document.status = 'failed'
        document.error = str(e)
        db.session.commit()
    finally:
        gc.collect()

def process_documents(file_paths, follow_up=False):
    """Process multiple documents and analyze them."""
    try:
        app.logger.info(f"Analyzing documents with Claude (follow_up: {follow_up})")
        
        # Check API key
        api_key = os.getenv('CLAUDE_API_KEY')
        if not api_key:
            raise ValueError("CLAUDE_API_KEY environment variable not set")
        app.logger.info(f"API Key length: {len(api_key)}")
        
        # Initialize Anthropic client
        client = anthropic.Anthropic(api_key=api_key, timeout=300)
        app.logger.info("Successfully initialized Anthropic client")
        
        # Process documents in smaller chunks
        MAX_TOKENS_PER_REQUEST = 3500
        all_results = []
        current_chunk = []
        current_chunk_tokens = 0
        
        # First pass: Calculate tokens for each document
        documents_with_tokens = []
        for file_path in file_paths:
            content = extract_text_from_document(file_path)
            tokens = count_tokens(content)
            app.logger.info(f"Document {os.path.basename(file_path)}: {tokens} tokens")
            documents_with_tokens.append((file_path, content, tokens))
        
        app.logger.info(f"Total tokens across all documents: {sum(doc[2] for doc in documents_with_tokens)}")
        
        # Second pass: Process documents in chunks
        for file_path, content, tokens in documents_with_tokens:
            # If this document would exceed token limit, process current chunk first
            if current_chunk_tokens + tokens > MAX_TOKENS_PER_REQUEST:
                app.logger.info(f"Processing chunk of {len(current_chunk)} documents ({current_chunk_tokens} tokens)")
                chunk_results = analyze_document_batch(current_chunk, client)
                all_results.extend(chunk_results)
                current_chunk = []
                current_chunk_tokens = 0
                
                # Force garbage collection
                gc.collect()
            
            # If single document is too large, split it
            if tokens > MAX_TOKENS_PER_REQUEST:
                app.logger.info(f"Large document detected: {os.path.basename(file_path)} ({tokens} tokens). Splitting into chunks.")
                doc_chunks = split_document_into_chunks(content, MAX_TOKENS_PER_REQUEST)
                for chunk_num, chunk in enumerate(doc_chunks, 1):
                    chunk_tokens = count_tokens(chunk)
                    app.logger.info(f"Processing chunk {chunk_num}/{len(doc_chunks)} of {os.path.basename(file_path)} ({chunk_tokens} tokens)")
                    chunk_result = analyze_document_batch([(file_path, chunk)], client)
                    all_results.extend(chunk_result)
                    gc.collect()
            else:
                current_chunk.append((file_path, content))
                current_chunk_tokens += tokens
        
        # Process any remaining documents
        if current_chunk:
            app.logger.info(f"Processing final chunk of {len(current_chunk)} documents ({current_chunk_tokens} tokens)")
            chunk_results = analyze_document_batch(current_chunk, client)
            all_results.extend(chunk_results)
        
        return all_results
        
    except Exception as e:
        app.logger.error(f"Error processing documents: {str(e)}")
        raise

def split_document_into_chunks(content, max_tokens):
    """Split a document into chunks that don't exceed max_tokens."""
    chunks = []
    sentences = content.split('. ')
    current_chunk = []
    current_chunk_tokens = 0
    
    for sentence in sentences:
        sentence_tokens = count_tokens(sentence)
        
        if current_chunk_tokens + sentence_tokens > max_tokens:
            # Join current chunk and add to chunks
            chunks.append('. '.join(current_chunk) + '.')
            current_chunk = [sentence]
            current_chunk_tokens = sentence_tokens
        else:
            current_chunk.append(sentence)
            current_chunk_tokens += sentence_tokens
    
    # Add any remaining content
    if current_chunk:
        chunks.append('. '.join(current_chunk) + '.')
    
    return chunks

def analyze_document_batch(document_batch, client):
    """Analyze a batch of documents."""
    results = []
    for file_path, content in document_batch:
        try:
            app.logger.info(f"Sending analysis request to Claude for {os.path.basename(file_path)}")
            
            message = client.messages.create(
                model="claude-3-opus-20240229",
                max_tokens=4000,
                temperature=0,
                system="You are a helpful assistant that analyzes legal documents.",
                messages=[
                    {
                        "role": "user",
                        "content": f"Please analyze this document and provide key information: {content}"
                    }
                ]
            )
            
            results.append({
                'file_path': file_path,
                'analysis': message.content
            })
            
            # Add a small delay between requests to prevent rate limiting
            time.sleep(1)
            
        except Exception as e:
            app.logger.error(f"Error analyzing document {os.path.basename(file_path)}: {str(e)}")
            results.append({
                'file_path': file_path,
                'error': str(e)
            })
    
    return results

@app.route('/')
def home():
    """Render home page and serve as health check endpoint."""
    try:
        # Quick DB check
        Property.query.limit(1).all()
        return render_template('index.html')
    except Exception as e:
        app.logger.error(f"Health check failed: {str(e)}")
        return "Service is starting up or experiencing issues", 503

@app.errorhandler(500)
def internal_server_error(e):
    """Handle internal server errors."""
    app.logger.error(f"Internal Server Error: {str(e)}")
    return jsonify({
        'error': 'Internal Server Error',
        'message': str(e),
        'suggestion': 'Please try again with a smaller file or fewer pages'
    }), 500

@app.errorhandler(413)
def request_entity_too_large(e):
    """Handle payload too large errors."""
    app.logger.error(f"Payload Too Large: {str(e)}")
    return jsonify({
        'error': 'File Too Large',
        'message': str(e),
        'suggestion': 'Please try uploading a smaller file (max 100MB)'
    }), 413

@app.route('/calculator_test')
def calculator():
    property_id = request.args.get('id')
    if property_id:
        property = Property.query.get_or_404(property_id)
        return render_template('calculator_test.html', property=property.to_dict())
    return render_template('calculator_test.html')

@app.route('/property/<int:property_id>')
def property_details(property_id):
    property = Property.query.get_or_404(property_id)
    return render_template('calculator_test.html', property=property.to_dict())

@app.route('/api/properties/<int:property_id>', methods=['DELETE'])
def delete_property(property_id):
    try:
        property = Property.query.get_or_404(property_id)
        db.session.delete(property)
        db.session.commit()
        return jsonify({'message': 'Property deleted successfully'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/properties/<int:property_id>', methods=['PUT'])
def update_property(property_id):
    try:
        data = request.get_json()
        
        # Update property fields
        property = Property.query.get_or_404(property_id)
        
        # Update property fields
        property.rightmove_url = data.get('rightmove_url')
        property.initial_cash = data.get('initial_cash')
        property.purchase_price = data.get('purchase_price')
        property.rooms = data.get('rooms')
        property.monthly_rent = data.get('monthly_rent')
        property.valuation_after = data.get('valuation_after')
        property.renovation_cost = data.get('renovation_cost')
        property.bridging_duration = data.get('bridging_duration')
        property.void_period = data.get('void_period')
        property.mortgage_ltv = data.get('mortgage_ltv')
        property.mortgage_rate = data.get('mortgage_rate')
        property.lender_fee = data.get('lender_fee')
        property.bridging_rate = data.get('bridging_rate')
        property.arrangement_rate = data.get('arrangement_rate')
        property.broker_rate = data.get('broker_rate')
        property.management_fee = data.get('management_fee')
        property.main_photo = data.get('main_photo')
        property.floorplan = data.get('floorplan')
        property.description = data.get('description')
        property.estate_agent = data.get('estate_agent')
        property.nearest_station = data.get('nearest_station')
        property.station_distance = data.get('station_distance')
        property.bedrooms = data.get('bedrooms')
        property.bathrooms = data.get('bathrooms')
        property.property_type = data.get('property_type')
        property.legal_pack_url = data.get('legal_pack_url')
        property.risk_level = data.get('risk_level')
        property.key_risks = data.get('key_risks')
        property.extra_fees = data.get('extra_fees')
        property.auction_date = datetime.strptime(data['auction_date'], '%Y-%m-%d').date() if data.get('auction_date') else None
        
        # Update viewing dates
        property.viewing_date_1 = datetime.fromisoformat(data['viewing_date_1'].replace('Z', '+00:00')) if data.get('viewing_date_1') else None
        property.viewing_date_2 = datetime.fromisoformat(data['viewing_date_2'].replace('Z', '+00:00')) if data.get('viewing_date_2') else None
        property.viewing_date_3 = datetime.fromisoformat(data['viewing_date_3'].replace('Z', '+00:00')) if data.get('viewing_date_3') else None
        property.viewing_date_4 = datetime.fromisoformat(data['viewing_date_4'].replace('Z', '+00:00')) if data.get('viewing_date_4') else None
        
        # Update calculated fields
        property.stamp_duty = data.get('stamp_duty')
        property.total_purchase_fees = data.get('total_purchase_fees')
        property.total_money_needed = data.get('total_money_needed')
        property.cash_left_in_deal = data.get('cash_left_in_deal')
        property.annual_profit = data.get('annual_profit')
        property.total_roi = data.get('total_roi')
        property.total_yield = data.get('total_yield')
        
        db.session.commit()
        return jsonify(property.to_dict()), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/properties/<int:property_id>/duplicate', methods=['POST'])
def duplicate_property(property_id):
    try:
        # Get the original property
        original = Property.query.get_or_404(property_id)
        
        # Create a new property with the same data
        new_property = Property(
            rightmove_url=original.rightmove_url,
            initial_cash=original.initial_cash,
            purchase_price=original.purchase_price,
            rooms=original.rooms,
            monthly_rent=original.monthly_rent,
            valuation_after=original.valuation_after,
            renovation_cost=original.renovation_cost,
            bridging_duration=original.bridging_duration,
            void_period=original.void_period,
            mortgage_ltv=original.mortgage_ltv,
            mortgage_rate=original.mortgage_rate,
            lender_fee=original.lender_fee,
            bridging_rate=original.bridging_rate,
            arrangement_rate=original.arrangement_rate,
            broker_rate=original.broker_rate,
            management_fee=original.management_fee,
            main_photo=original.main_photo,
            floorplan=original.floorplan,
            description=original.description,
            key_features=original.key_features,
            address=f"{original.address} (Copy)",
            is_auction=original.is_auction,
            estate_agent=original.estate_agent,
            nearest_station=original.nearest_station,
            station_distance=original.station_distance,
            bedrooms=original.bedrooms,
            bathrooms=original.bathrooms,
            property_type=original.property_type,
            legal_pack_url=original.legal_pack_url,
            legal_pack_available=original.legal_pack_available,
            risk_level=original.risk_level,
            key_risks=original.key_risks,
            extra_fees=original.extra_fees,
            auction_date=original.auction_date,
            stamp_duty=original.stamp_duty,
            total_purchase_fees=original.total_purchase_fees,
            total_money_needed=original.total_money_needed,
            cash_left_in_deal=original.cash_left_in_deal,
            annual_profit=original.annual_profit,
            total_roi=original.total_roi,
            total_yield=original.total_yield,
            legal_pack_analysis=original.legal_pack_analysis,
            legal_pack_qa_history=original.legal_pack_qa_history,
            legal_pack_documents=original.legal_pack_documents,
            legal_pack_summary_pdf=original.legal_pack_summary_pdf,
            legal_pack_analyzed_at=original.legal_pack_analyzed_at,
            legal_pack_session_id=original.legal_pack_session_id
        )
        
        db.session.add(new_property)
        db.session.commit()
        
        return jsonify({'message': 'Property duplicated successfully', 'id': new_property.id})
    except Exception as e:
        return jsonify({'error': f'Failed to duplicate property: {str(e)}'}), 500

@app.route('/api/properties/<int:property_id>', methods=['GET'])
def get_property(property_id):
    try:
        property = Property.query.get_or_404(property_id)
        return jsonify(property.to_dict())
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/properties', methods=['GET'])
def get_properties():
    try:
        properties = Property.query.order_by(Property.created_at.desc()).all()
        return jsonify([p.to_dict() for p in properties])
    except Exception as e:
        print("Error fetching properties:", str(e))
        return jsonify({'error': 'Failed to fetch properties', 'details': str(e)}), 500

@app.route('/api/properties', methods=['POST'])
def create_property():
    try:
        data = request.get_json()
        
        # Handle auction_date - set to None if empty string or invalid
        auction_date = None
        if data.get('auction_date'):
            if isinstance(data['auction_date'], str) and data['auction_date'].strip():
                try:
                    auction_date = datetime.strptime(data['auction_date'], '%Y-%m-%d').date()
                except ValueError:
                    auction_date = None

        # Convert key_features to JSON string if it's a list
        key_features = data.get('key_features', [])
        if isinstance(key_features, list):
            key_features = json.dumps(key_features)
        elif key_features is None:
            key_features = json.dumps([])

        property = Property(
            rightmove_url=data.get('rightmove_url'),
            initial_cash=data.get('initial_cash'),
            purchase_price=data.get('purchase_price'),
            rooms=data.get('rooms'),
            monthly_rent=data.get('monthly_rent'),
            valuation_after=data.get('valuation_after'),
            renovation_cost=data.get('renovation_cost'),
            bridging_duration=data.get('bridging_duration'),
            void_period=data.get('void_period'),
            mortgage_ltv=data.get('mortgage_ltv'),
            mortgage_rate=data.get('mortgage_rate'),
            lender_fee=data.get('lender_fee'),
            bridging_rate=data.get('bridging_rate'),
            arrangement_rate=data.get('arrangement_rate'),
            broker_rate=data.get('broker_rate'),
            management_fee=data.get('management_fee'),
            created_at=datetime.now(),
            main_photo=data.get('main_photo'),
            floorplan=data.get('floorplan'),
            description=data.get('description'),
            key_features=key_features,
            address=data.get('address'),
            is_auction=data.get('is_auction', False),
            estate_agent=data.get('estate_agent'),
            nearest_station=data.get('nearest_station'),
            station_distance=data.get('station_distance'),
            bedrooms=data.get('bedrooms'),
            bathrooms=data.get('bathrooms'),
            property_type=data.get('property_type'),
            legal_pack_url=data.get('legal_pack_url'),
            legal_pack_available=data.get('legal_pack_available', False),
            risk_level=data.get('risk_level'),
            key_risks=data.get('key_risks'),
            extra_fees=data.get('extra_fees', 0),
            auction_date=auction_date,
            stamp_duty=data.get('stamp_duty'),
            total_purchase_fees=data.get('total_purchase_fees'),
            total_money_needed=data.get('total_money_needed'),
            cash_left_in_deal=data.get('cash_left_in_deal'),
            annual_profit=data.get('annual_profit'),
            total_roi=data.get('total_roi'),
            total_yield=data.get('total_yield'),
            legal_pack_analysis=data.get('legal_pack_analysis'),
            legal_pack_qa_history=data.get('legal_pack_qa_history'),
            legal_pack_documents=data.get('legal_pack_documents'),
            legal_pack_summary_pdf=data.get('legal_pack_summary_pdf'),
            legal_pack_analyzed_at=data.get('legal_pack_analyzed_at'),
            legal_pack_session_id=data.get('legal_pack_session_id'),
            viewing_date_1=datetime.fromisoformat(data['viewing_date_1'].replace('Z', '+00:00')) if data.get('viewing_date_1') else None,
            viewing_date_2=datetime.fromisoformat(data['viewing_date_2'].replace('Z', '+00:00')) if data.get('viewing_date_2') else None,
            viewing_date_3=datetime.fromisoformat(data['viewing_date_3'].replace('Z', '+00:00')) if data.get('viewing_date_3') else None,
            viewing_date_4=datetime.fromisoformat(data['viewing_date_4'].replace('Z', '+00:00')) if data.get('viewing_date_4') else None,
        )
        
        db.session.add(property)
        db.session.commit()
        
        return jsonify(property.to_dict()), 201
    except Exception as e:
        db.session.rollback()
        print("Error details:", str(e))
        return jsonify({'error': 'Failed to create property', 'details': str(e)}), 400

@app.route('/scrape-property', methods=['POST'])
def scrape_property():
    url = request.json.get('url')
    if not url:
        return jsonify({'error': 'URL is required'}), 400

    scraper = PropertyScraper()
    try:
        # Run the async scraping in the event loop
        result = asyncio.run(scraper.scrape_rightmove(url))
        if result:
            return jsonify(result), 200
        return jsonify({'error': 'Failed to scrape property data'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 400

@app.route('/toggle_auction/<int:property_id>', methods=['POST'])
def toggle_auction(property_id):
    try:
        data = request.get_json()
        property = Property.query.get(property_id)
        property.is_auction = data['is_auction']
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 400

@app.route('/legal-pack-analyzer/<int:property_id>')
def legal_pack_analyzer(property_id):
    """Render the legal pack analyzer page for a specific property."""
    # Get the property details if needed
    property_data = Property.query.get(property_id)
    return render_template('legal_pack_analyzer.html', property=property_data, property_id=property_id)

@app.route('/analyze-legal-pack', methods=['POST'])
def analyze_legal_pack():
    try:
        app.logger.info("Starting legal pack analysis...")
        
        # Check if we have a single file (ZIP) or multiple files
        if 'file' in request.files:
            file = request.files['file']
            if not file or file.filename == '':
                app.logger.error("No file selected")
                return jsonify({'error': 'No file selected'}), 400
                
            if not file.filename.lower().endswith('.zip'):
                app.logger.error("File must be a ZIP archive")
                return jsonify({'error': 'File must be a ZIP archive'}), 400
                
            property_id = request.form.get('property_id')
            if not property_id:
                app.logger.error("Property ID is missing")
                return jsonify({'error': 'Property ID is required'}), 400
                
            app.logger.info(f"Processing ZIP file for property ID: {property_id}")
            
            # Create a temporary directory for processing
            with tempfile.TemporaryDirectory() as temp_dir:
                app.logger.info(f"Created temporary directory: {temp_dir}")
                
                # Save and process the ZIP file
                zip_path = os.path.join(temp_dir, secure_filename(file.filename))
                file.save(zip_path)
                
                try:
                    processed_files, failed_files, processing_summary = process_zip_file(zip_path)
                    if not processed_files:
                        app.logger.error("No valid files found in ZIP archive")
                        return jsonify({
                            'error': 'No valid files found in ZIP archive',
                            'processing_summary': processing_summary
                        }), 400
                        
                    # Combine all text from processed files
                    all_text = [doc['content'] for doc in processed_files]
                    combined_text = '\n\n'.join(all_text)
                    total_chars = len(combined_text)
                    total_words = len(combined_text.split())
                    
                    app.logger.info("Document processing completed:")
                    app.logger.info(f"- Total files processed: {len(processed_files)}")
                    app.logger.info(f"- Total characters: {total_chars}")
                    app.logger.info(f"- Total words: {total_words}")
                    
                    # Create analysis record
                    try:
                        # Store the extracted text in Analysis table
                        analysis = Analysis(
                            property_id=property_id,
                            content=combined_text,
                            timestamp=datetime.utcnow()
                        )
                        db.session.add(analysis)
                        db.session.commit()
                        app.logger.info(f"Analysis saved to database with ID: {analysis.id}")

                        # Perform Claude analysis
                        app.logger.info("Starting Claude analysis...")
                        session_id = str(uuid.uuid4())
                        analysis_result = analyze_with_claude(
                            documents_content=processed_files,
                            processing_summary=processing_summary
                        )

                        if analysis_result:
                            # Update property with analysis results
                            property_record = Property.query.get(property_id)
                            if property_record:
                                property_record.legal_pack_analysis = analysis_result
                                property_record.legal_pack_analyzed_at = datetime.utcnow()
                                property_record.legal_pack_session_id = session_id
                                db.session.commit()
                                app.logger.info("Analysis results saved to property record")

                        # Calculate token usage for each document
                        token_usage = {
                            'total_tokens': sum(doc['tokens'] for doc in processed_files),
                            'documents': [
                                {
                                    'name': doc['name'],
                                    'tokens': doc['tokens']
                                }
                                for doc in processed_files
                            ]
                        }

                        return jsonify({
                            'message': 'Analysis completed successfully',
                            'session_id': session_id,
                            'analysis_id': analysis.id,
                            'analysis': analysis_result,
                            'stats': {
                                'total_files': len(processed_files),
                                'total_characters': total_chars,
                                'total_words': total_words,
                                'failed_files': len(failed_files)
                            },
                            'token_usage': token_usage,
                            'processing_summary': processing_summary
                        })
                    except Exception as e:
                        app.logger.error(f"Error saving to database: {str(e)}")
                        return jsonify({
                            'error': f'Error saving analysis: {str(e)}',
                            'processing_summary': processing_summary
                        }), 500
                except Exception as e:
                    app.logger.error(f"Error processing ZIP file: {str(e)}")
                    return jsonify({
                        'error': f'Error processing ZIP file: {str(e)}'
                    }), 500
        else:
            app.logger.error("No file uploaded")
            return jsonify({'error': 'No file uploaded'}), 400
    except Exception as e:
        app.logger.error(f"Error in analyze_legal_pack: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/property/ask_followup', methods=['POST'])
def ask_followup():
    """Handle follow-up questions about the legal pack."""
    try:
        data = request.get_json()
        question = data.get('question')
        session_id = data.get('session_id')
        property_id = data.get('property_id')
        
        logger.info(f"Received follow-up question request: {data}")
        
        if not all([question, session_id, property_id]):
            missing = []
            if not question: missing.append('question')
            if not session_id: missing.append('session_id')
            if not property_id: missing.append('property_id')
            error_msg = f"Missing required information: {', '.join(missing)}"
            logger.error(error_msg)
            return jsonify({
                'error': error_msg,
                'suggestion': 'Please provide all required information'
            }), 400
        
        # Load the documents and previous analysis from the property
        try:
            property = Property.query.get(property_id)
            if not property:
                return jsonify({
                    'error': 'Property not found',
                    'suggestion': 'Please check the property ID'
                }), 404

            if not property.legal_pack_analysis or not property.legal_pack_documents:
                return jsonify({
                    'error': 'No legal pack analysis or documents found',
                    'suggestion': 'Please analyze the legal pack first'
                }), 404
                
            # Get QA history and documents from the property
            qa_history = json.loads(property.legal_pack_qa_history) if property.legal_pack_qa_history else []
            
            # Get the original documents from the Analysis table
            analysis = Analysis.query.filter_by(property_id=property_id).order_by(Analysis.timestamp.desc()).first()
            if not analysis:
                return jsonify({
                    'error': 'No analysis record found',
                    'suggestion': 'Please analyze the legal pack first'
                }), 404
                
            # Format the documents for Claude
            documents = [{
                'name': f'Document {i+1}',
                'content': content.strip()
            } for i, content in enumerate(analysis.content.split('\n\n')) if content.strip()]
            
            # Get answer from Claude
            try:
                result = analyze_with_claude(
                    documents_content=documents,  # Pass the formatted documents
                    follow_up_question=question,
                    initial_analysis=property.legal_pack_analysis,
                    qa_history=qa_history
                )
                logger.info("Successfully got answer from Claude")
                
                if not result:
                    raise ValueError("Empty response from Claude")
                    
            except Exception as e:
                error_msg = f"Failed to get answer from Claude: {str(e)}"
                logger.error(error_msg)
                return jsonify({
                    'error': error_msg,
                    'suggestion': 'Please try again or contact support if the problem persists'
                }), 500
            
            # Update QA history
            qa_history.append({
                'question': question,
                'answer': result,
                'timestamp': datetime.utcnow().isoformat()
            })
            property.legal_pack_qa_history = json.dumps(qa_history)
            db.session.commit()
            logger.info("Updated QA history")
            
            # Calculate token usage
            token_usage = {
                'total_tokens': sum(count_tokens(doc['content']) for doc in documents),
                'documents': [
                    {
                        'name': doc['name'],
                        'tokens': count_tokens(doc['content'])
                    }
                    for doc in documents
                ]
            }
            
            return jsonify({
                'answer': result,
                'token_usage': token_usage,
                'qa_history': qa_history
            })
                
        except Exception as e:
            error_msg = f"Failed to process follow-up question: {str(e)}"
            logger.error(error_msg)
            return jsonify({
                'error': error_msg,
                'suggestion': 'Please try again or contact support if the problem persists'
            }), 500
    except Exception as e:
        error_msg = f"Unexpected error in ask_followup: {str(e)}"
        logger.error(error_msg)
        return jsonify({
            'error': error_msg,
            'suggestion': 'Please try again or contact support if the problem persists'
        }), 500

@app.route('/system-check', methods=['GET'])
def system_check():
    """Check system dependencies and configuration."""
    try:
        # Check environment variables
        env_vars = {
            'CLAUDE_API_KEY': bool(os.getenv('CLAUDE_API_KEY')),
            'DATABASE_URL': bool(os.getenv('DATABASE_URL'))
        }
        
        status = {
            'tesseract': False,
            'libreoffice': False,
            'environment': env_vars
        }
        
        # Check tesseract
        tesseract_path = shutil.which('tesseract')
        status['tesseract'] = bool(tesseract_path)
        
        # Check libreoffice
        soffice_path = shutil.which('soffice')
        status['libreoffice'] = bool(soffice_path)
        
        return jsonify(status)
    except Exception as e:
        app.logger.error(f"Error in system check: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/test-property-details')
def test_property_details():
    # Sample property data for testing
    test_property = {
        'main_photo': 'https://media.rightmove.co.uk/dir/crop/10:9-16:9/108k/107051/128095236/107051_11261955_IMG_00_0000_max_476x317.jpeg',
        'floorplan': 'https://media.rightmove.co.uk/dir/108k/107051/128095236/107051_11261955_FLP_00_0000_max_600x600.jpeg',
        'address': '123 Test Street, London',
        'description': 'A beautiful test property with modern amenities and great location.',
        'bedrooms': 3,
        'bathrooms': 2,
        'property_type': 'Semi-Detached',
        'station_distance': '0.5',
        'key_features': [
            'Modern Kitchen',
            'Large Garden',
            'Recently Renovated',
            'Close to Schools'
        ],
        'purchase_price': 250000,
        'legal_fees': 1500,
        'survey_costs': 800,
        'refurb_costs': 15000,
        'contingency': 5000,
        'bridging_loan_amount': 200000,
        'bridging_loan_term': 6,
        'bridging_loan_rate': 0.89,
        'mortgage_amount': 187500,
        'mortgage_term': 25,
        'mortgage_rate': 4.5,
        'rental_income': 1500
    }
    return render_template('property_details.html', property=test_property)

@app.route('/document_status/<int:doc_id>')
def document_status(doc_id):
    """Get document processing status."""
    try:
        document = Document.query.get_or_404(doc_id)
        return jsonify({
            'status': document.status,
            'processed_pages': document.processed_pages,
            'total_pages': document.total_pages,
            'error': document.error,
            'text_content': document.text_content if document.status == 'completed' else None
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def log_memory_usage(context=""):
    """Log current memory usage."""
    try:
        import psutil
        process = psutil.Process()
        memory_info = process.memory_info()
        logger.info(f"Memory usage ({context}): {memory_info.rss / 1024 / 1024:.2f} MB")
    except ImportError:
        logger.warning("psutil not installed - cannot monitor memory usage")

@app.before_request
def before_request():
    """Log memory usage before each request."""
    log_memory_usage(f"Before request: {request.path}")

@app.after_request
def after_request(response):
    """Log memory usage after each request."""
    log_memory_usage(f"After request: {request.path}")
    return response

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True, port=5004)
else:
    # This ensures tables are created when running on Render
    with app.app_context():
        db.create_all()
